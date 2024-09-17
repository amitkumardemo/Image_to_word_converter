import streamlit as st
from PIL import Image
import pytesseract
from docx import Document
import io

# Set the Tesseract executable path (if necessary for Windows)
# Uncomment and set the path if running on Windows
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Set page configuration
st.set_page_config(page_title="Image to Word Converter", layout="wide")

# Add CSS for custom styles
st.markdown("""
<style>
/* Navbar styles */
.navbar {
    background-color: #333;
    overflow: hidden;
}
.navbar a {
    float: left;
    display: block;
    color: #f2f2f2;
    text-align: center;
    padding: 14px 16px;
    text-decoration: none;
}
.navbar a:hover {
    background-color: #ddd;
    color: black;
}

/* Footer styles */
.footer {
    background-color: #333;
    color: white;
    text-align: center;
    padding: 10px;
    position: fixed;
    width: 100%;
    bottom: 0;
}
.footer a {
    color: white;
    text-decoration: none;
    margin: 0 10px;
}
.footer a:hover {
    text-decoration: underline;
}
</style>
""", unsafe_allow_html=True)

# Logo
st.image("jb.png", width=250)  # Replace with your logo URL

# Navbar
st.markdown("""
<div class="navbar">
  <a href="#Home">Home</a>
  <a href="#About">About</a>
  <a href="https://techiehelpt.netlify.app/">BackToWebsite</a> <!-- Replace with your actual website URL -->
</div>
""", unsafe_allow_html=True)

# Add navigation to different sections
menu = ["Home", "About"]
choice = st.sidebar.selectbox("Navigate", menu)

if choice == "Home":
    # Home Page
    st.title("Image to Word Converter")
    st.write("Upload an image to convert its text into a Word document.")

    # File uploader for image input
    uploaded_file = st.file_uploader("Choose an image...", type=["jpg", "png", "jpeg"])

    if uploaded_file is not None:
        # Open the uploaded image
        image = Image.open(uploaded_file)

        # Display the file name instead of the full image
        st.write(f"File name: **{uploaded_file.name}**")

        # Extract text from the image using pytesseract
        st.write("Extracting text from image...")
        try:
            text = pytesseract.image_to_string(image)

            # Display the extracted text in a text area
            st.text_area("Extracted Text", text, height=300)

            # Create a Word document
            doc = Document()
            doc.add_heading('Extracted Text', level=1)
            doc.add_paragraph(text)
            
            # Save the document to a BytesIO stream
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Download button for the Word document
            st.download_button(
                label="📥 Download Word Document",
                data=buffer,
                file_name="extracted_text.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        except pytesseract.TesseractNotFoundError:
            st.error("Tesseract OCR not found. Please ensure Tesseract is installed and the path is set correctly.")

elif choice == "About":
    # About Page
    st.title("About Image to Word Converter")
    st.write("""
    This tool allows you to upload an image and convert the text contained in the image into a Word document.
    The project uses the Python library `pytesseract` for OCR and `python-docx` for creating Word documents.

    **Features**:
    - Upload an image in JPG, PNG, or JPEG format.
    - Extract text from the image.
    - Download the extracted text as a Word document.

    **Technology Stack**:
    - Streamlit (for building the web app)
    - pytesseract (for OCR)
    - Pillow (for image handling)
    - python-docx (for creating Word documents)
    """)

# Footer
st.markdown("""
<div class="footer">
    <p>© 2024 Image to Word Converter | TechieHelp</p>
    <a href="https://www.linkedin.com/in/techiehelp" target="_blank">LinkedIn</a>
    <a href="https://www.twitter.com/techiehelp" target="_blank">Twitter</a>
    <a href="https://www.instagram.com/techiehelp2" target="_blank">Instagram</a>
</div>
""", unsafe_allow_html=True)
